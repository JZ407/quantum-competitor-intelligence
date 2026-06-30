"""Generate QuEra roadmap analysis Word report."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def p(text):
    return doc.add_paragraph(text)

def b(text):
    return doc.add_paragraph(text, style='List Bullet')

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        t.rows[0].cells[i].text = h_text
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        r = t.add_row()
        for i, cell in enumerate(row):
            r.cells[i].text = str(cell)
    return t

# ═══ Title ═══
title = doc.add_heading('QuEra 2028 年容错量子计算路线图可信度分析', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('2026年6月26日\n').font.size = Pt(10)
meta.add_run('核心问题：QuEra 声称 2028 年实现 1000+ 逻辑量子比特、10⁻⁹ 保真度（99.9999999%），\n其底气从何而来？是否可信？').font.size = Pt(10)

# ═══ 一 ═══
h('一、核心结论', 1)
p('QuEra 的 2028 年路线图并非凭空画饼，而是建立在"三根支柱"之上：')
b('架构创新（Transversal STAR）：解决了传统容错量子计算的两个最大资源瓶颈——魔态蒸馏和晶格手术路由')
b('中性原子平台天然优势：全连通、常温运行、单模块扩展至数千原子，工程路径更短')
b('持续的论文验证：8 篇 Nature/PRL 论文逐层验证了从逻辑比特、横向门到超高率纠错的每一个技术环节')
p('评估：路线图激进但可信（7.5/10）。主要风险在于工程放大（从实验室 3,000 原子到商用 10,000+ 原子的良率）和 2 年交付周期对执行力的要求。')

# ═══ 二 ═══
h('二、技术底气拆解', 1)

h('2.1 架构创新：Transversal STAR', 2)
p('传统容错量子计算 ~90% 的资源消耗在魔态蒸馏上。QuEra 与 Los Alamos 合作发表的 Transversal STAR 架构（2026年5月 PRX Quantum）绕过了这一瓶颈：')

table(
    ['指标', '传统 FTQC', 'Transversal STAR', '优势'],
    [
        ['魔态制备', '蒸馏工厂占 ~90% 物理比特', '横向注入 + post-selection', '物理比特需求减半'],
        ['Clifford 门', '晶格手术 ~d 轮校验子', '原子 shuttling ~1 轮', '速度快 ~250 倍'],
        ['编码效率', 'surface code ~1:81', 'qLDPC 可达 1:2', '物理比特需求再降 5 倍'],
    ]
)

p('关键数字：使用 qLDPC 编码后，仅需 1,500-3,000 物理比特即可达到 megaquop 级别（100万次可靠逻辑操作）。常规架构需要 10万+ 物理比特。')

h('2.2 中性原子平台 vs 竞品', 2)

table(
    ['参数', '中性原子 (QuEra)', '超导 (IBM/Google)', '离子阱 (Quantinuum)'],
    [
        ['运行温度', '常温（真空腔）', '~10 mK（稀释制冷机）', '常温'],
        ['量子比特连接', '全连通（光镊移动）', '近邻（固定芯片）', '全连通'],
        ['扩展方式', '单模块加原子', '多芯片互联', '多阱互联'],
        ['最大阵列（已展示）', '3,000+ 原子', '1,121 量子比特', '56 量子比特'],
        ['门保真度（二比特）', '~99.92%', '~99.85%', '~99.8%'],
    ]
)

p('中性原子的"全连通 + 可动态移动"是 Transversal STAR 架构成立的物理基础。')

h('2.3 论文验证时间线', 2)
p('论文产出与路线图升级存在清晰的因果链。Google/软银 $230M 不是在"先给钱再研发"，而是看到论文验证后才投的。')

table(
    ['时间', '里程碑', '详情'],
    [
        ['2023.12', '48 逻辑比特展示', 'Nature: Logical quantum processor — Physics World 年度突破'],
        ['2024.03', '恒开销容错理论', 'Nature Physics: Constant-Overhead FTQC'],
        ['2024.04', '关联解码横向逻辑门', 'PRL: Correlated decoding of logical algorithms'],
        ['2024.12', '逻辑魔态蒸馏实验', 'Nature: Magic State Distillation 验证'],
        ['2025.02', '\U0001f4b0 $230M 融资', 'Google/软银领投——论文验证吸引资本'],
        ['2025.06', '6 篇密集产出', '横向架构 + qLDPC + 魔态蒸馏 + 编译优化'],
        ['2025.09', 'STAR 架构预印本', 'Nature + arXiv: Transversal STAR'],
        ['2026.04', 'Tsim + 超高率 QEC', '10⁻¹³ 错误率验证'],
        ['2026.05', 'STAR 正式发表', 'PRX Quantum: Transversal Architecture'],
        ['2026.06', '\U0001f3af Libra 2028 宣布', '256逻辑比特,10⁻⁶ → 1000+,10⁻⁹'],
    ]
)

# ═══ 三 ═══
h('三、两阶段路线图', 1)

table(
    ['指标', 'Libra (2028)', 'Next Gen (2028-29)'],
    [
        ['逻辑量子比特', '256', '1,000+'],
        ['物理量子比特', '10,000-15,000', '20,000+'],
        ['逻辑错误率', '10⁻⁶', '10⁻⁹'],
        ['可靠逻辑操作', '100万次 (megaquop)', '10亿次 (gigaquop)'],
        ['部署平台', 'AWS Braket', 'AWS Braket'],
        ['典型应用', '量子模拟', '量子化学/高能物理/聚变'],
    ]
)

# ═══ 四 ═══
h('四、风险评估', 1)

table(
    ['风险类别', '等级', '描述'],
    [
        ['工程风险', '\U0001f7e1 中等', '3,000 → 10,000+ 原子良率；实时原子重载复杂度；GPU实时解码延迟'],
        ['时间风险', '\U0001f534 中高', '2年从论文到商用；AWS集成需额外软件栈开发'],
        ['人才风险', '\U0001f7e2 中低', '核心依赖四位学术创始人；量子工程人才稀缺'],
        ['竞争风险', '\U0001f7e1 中等', 'IBM 2029年FTQC（仅晚1年）；Pasqal欧陆竞争；Atom+Quantinuum合并'],
    ]
)

# ═══ 五 ═══
h('五、综合判断', 1)
p('QuEra 的 2028 路线图可信度：较高（7.5/10）')

h('加分项', 2)
b('技术路线有清晰的理论和实验基础，非凭空想象')
b('论文-融资-产品节奏高度协调，符合深科技公司发展规律')
b('中性原子平台在扩展性和工程化方面确实有独特优势')

h('减分项', 2)
b('2 年交付周期对硬件公司而言极短')
b('从 3,000 到 10,000+ 原子的良率跃升尚未公开展示')
b('竞争对手的追赶速度可能压缩先发优势窗口')

p('\n对领导层的建议：密切关注 QuEra 在 2026年下半年到2027年初的实际硬件里程碑（如 Gemini 系统的逻辑比特保真度数据），这将是判断 2028年能否兑现的关键先行指标。')

# Footer
doc.add_paragraph('')
fp = doc.add_paragraph('本报告基于 QuEra 官网公开信息、89篇论文数据、PRX Quantum/Nature 等期刊发表记录、以及 2026年6月 QuEra-AWS 联合公告综合分析。')
fp.runs[0].font.size = Pt(9)
fp.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.save('D:/Claude_code/competitor_profiles/quera_roadmap_analysis.docx')
print('Saved: quera_roadmap_analysis.docx')
